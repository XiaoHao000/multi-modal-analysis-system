"""
Word 文档处理器: .docx / .doc 文本提取。

流程:
  .docx → python-docx 提取段落文本
  .doc  → python-docx 尝试打开（兼容新版），失败则返回空

优雅降级: 解析失败返回空字符串，不阻塞主流程。
"""

import base64
import io
from typing import List

from utils.logger import logger


def _decode_to_bytes(data_b64: str) -> bytes:
    return base64.b64decode(data_b64)


def _extract_docx(file_bytes: bytes) -> str:
    """从 .docx 文件中提取所有段落文本。"""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    # 也提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                paragraphs.append(row_text)
    return "\n\n".join(paragraphs)


def _extract_doc(file_bytes: bytes) -> str:
    """从 .doc 文件中提取文本（尽可能兼容）。"""
    try:
        # 新版 python-docx 可能无法打开旧 .doc，尝试一下
        return _extract_docx(file_bytes)
    except Exception:
        pass
    # 尝试用 antiword/textract 等外部工具
    logger.warning(".doc 文件无法解析，请转换为 .docx 格式后上传")
    return ""


async def process_word_files(files: List[dict], user_query: str = "") -> str:
    """处理 Word 文件：提取文本内容。

    Args:
        files: Word 文件列表 [{"mime_type": "application/vnd...document", "data": "base64...", "filename": "report.docx"}]
        user_query: 用户原始问题（保留接口一致，当前未使用）

    Returns:
        word_text: 提取的纯文本字符串
    """
    if not files:
        return ""

    all_text = []

    for entry in files:
        data = entry.get("data", "")
        filename = entry.get("filename", "unknown")
        mime = entry.get("mime_type", "")
        if not data:
            continue

        try:
            file_bytes = _decode_to_bytes(data)
        except Exception as e:
            logger.warning(f"Word base64 解码失败 ({filename}): {e}")
            continue

        try:
            if "openxmlformats" in mime or filename.lower().endswith(".docx"):
                text = _extract_docx(file_bytes)
            else:
                text = _extract_doc(file_bytes)
        except Exception as e:
            logger.warning(f"Word 提取失败 ({filename}): {e}")
            text = ""

        if text:
            all_text.append(f"--- {filename} ---\n{text}")
        logger.info(
            f"Word 处理完成: {filename} → {len(text)} 字"
        )

    return "\n\n".join(all_text)
